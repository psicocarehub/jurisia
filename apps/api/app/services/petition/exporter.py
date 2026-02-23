"""
Export petitions to PDF and DOCX formats.

Uses xhtml2pdf for PDF and python-docx for DOCX generation.
Both include CNJ Res. 615/2025 disclaimer in the footer when AI-generated.
"""

import io
import re
from typing import Optional

from app.services.petition.formatter import AI_DISCLAIMER


def _build_html(content: str, title: str, ai_generated: bool = False) -> str:
    """Wrap petition content in a full HTML document with ABNT styling."""
    disclaimer_block = ""
    if ai_generated:
        disclaimer_block = f"""
        <div style="border-top: 1px solid #ccc; padding-top: 10pt; margin-top: 30pt;
                    font-size: 9pt; color: #666; text-align: center;">
            <strong>Aviso:</strong> {AI_DISCLAIMER}
        </div>
        """

    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="utf-8"/>
    <style>
        @page {{
            size: A4;
            margin: 3cm 2cm 2cm 3cm;
        }}
        body {{
            font-family: 'Times New Roman', Times, serif;
            font-size: 12pt;
            line-height: 1.5;
            color: #000;
            text-align: justify;
        }}
        h1 {{
            font-size: 16pt;
            font-weight: bold;
            text-align: center;
            text-transform: uppercase;
            margin-bottom: 24pt;
        }}
        h2 {{
            font-size: 14pt;
            font-weight: bold;
            text-align: center;
            text-transform: uppercase;
            margin-top: 18pt;
            margin-bottom: 12pt;
        }}
        p {{
            text-indent: 1.25cm;
            margin-bottom: 6pt;
        }}
        blockquote {{
            font-size: 10pt;
            line-height: 1.0;
            margin-left: 4cm;
            margin-top: 12pt;
            margin-bottom: 12pt;
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    {content}
    {disclaimer_block}
</body>
</html>"""


def export_pdf(
    content: str,
    title: str = "Peticao",
    ai_generated: bool = False,
    metadata: Optional[dict] = None,
) -> bytes:
    """Export petition content to PDF bytes."""
    from xhtml2pdf import pisa

    html = _build_html(content, title, ai_generated)
    buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(io.StringIO(html), dest=buffer)

    if pisa_status.err:
        raise RuntimeError(f"PDF generation failed with {pisa_status.err} errors")

    return buffer.getvalue()


def export_docx(
    content: str,
    title: str = "Peticao",
    ai_generated: bool = False,
    metadata: Optional[dict] = None,
) -> bytes:
    """Export petition content to DOCX bytes."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)

    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)

    clean_text = re.sub(r'<[^>]+>', '\n', content)
    paragraphs = [p.strip() for p in clean_text.split('\n') if p.strip()]

    for para_text in paragraphs:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if ai_generated:
        doc.add_paragraph()
        disclaimer = doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = disclaimer.add_run(f"Aviso: {AI_DISCLAIMER}")
        run.font.size = Pt(9)
        run.font.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
